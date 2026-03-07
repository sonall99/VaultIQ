"""
VaultIQ — Document Parser
Converts uploaded files (PDF, DOCX, TXT, CSV) to plain text
for ingestion into the RAG pipeline.
"""

import io
from typing import Union


def parse_file_bytes(file_bytes: bytes, extension: str) -> str:
    """
    Parse a file from raw bytes to plain text.

    Supported formats:
      - .txt / .md  → read directly
      - .pdf        → extract with PyMuPDF (fitz)
      - .docx       → extract with python-docx
      - .csv        → read as plain text rows

    Args:
        file_bytes: raw file bytes from upload
        extension:  file extension without dot, e.g. "pdf"

    Returns:
        Extracted plain text string
    """
    ext = extension.lower().strip(".")

    if ext in ("txt", "md"):
        return _parse_txt(file_bytes)
    elif ext == "pdf":
        return _parse_pdf(file_bytes)
    elif ext == "docx":
        return _parse_docx(file_bytes)
    elif ext == "csv":
        return _parse_csv(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Supported: txt, pdf, docx, csv")


# ── Parsers ───────────────────────────────────────────────

def _parse_txt(file_bytes: bytes) -> str:
    """Plain text — decode UTF-8, fallback to latin-1"""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1")


def _parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF (fitz).
    Handles multi-page PDFs, preserves paragraph structure.
    Install: pip install PyMuPDF
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")  # plain text extraction
            if text.strip():
                pages.append(f"[Page {page_num + 1}]\n{text.strip()}")

        doc.close()
        return "\n\n".join(pages)

    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install PyMuPDF")
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def _parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX using python-docx.
    Extracts paragraphs and table cells.
    Install: pip install python-docx
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        sections = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                sections.append(text)

        # Extract table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    sections.append(row_text)

        return "\n\n".join(sections)

    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def _parse_csv(file_bytes: bytes) -> str:
    """
    Convert CSV to plain text rows.
    Each row becomes a readable line.
    """
    try:
        import csv

        text = file_bytes.decode("utf-8")
        reader = csv.reader(io.StringIO(text))
        rows = [", ".join(row) for row in reader if any(cell.strip() for cell in row)]
        return "\n".join(rows)

    except Exception as e:
        raise ValueError(f"Failed to parse CSV: {str(e)}")


# ── Utility ───────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Clean extracted text:
    - Remove excessive whitespace
    - Remove non-printable characters
    - Normalise line breaks
    """
    import re

    # Remove non-printable chars except newlines/tabs
    text = re.sub(r"[^\x20-\x7E\n\t]", " ", text)
    # Collapse multiple spaces
    text = re.sub(r"[ \t]+", " ", text)
    # Collapse 3+ newlines to 2
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()
